import streamlit as st
import math
from tqdm import tqdm
import traceback
import os
import sys
import concurrent.futures
import nltk
from spire.doc import *
from spire.doc.common import *
from docx import Document as doc
nltk.download('punkt')

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

# Get the absolute path to the main directory (my_project)
main_directory = os.path.abspath(os.path.dirname(__file__))
# Add main directory and its subdirectories to sys.path
sys.path.append(main_directory)
sys.path.append("/mount/src/carlat_qa_live_update/src")
sys.path.append("/mount/src/carlat_qa_live_update")

from backend.setup_vectorstore import *
from backend.utility_functions import *
from backend.utility_functions import find_redundant_quotes
from backend.retreive import *
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"



def update_file_params():
    # update the upload control parameters
    st.session_state.file_uploaded = False
    st.session_state.file_content = ""
    file = None


def process_quote(topic, quote,  custom_qa_prompts):
    #  process each quote
    qa = get_qa_for_quote(quote, topic, custom_qa_prompts)
    qa_json = json.loads(qa)

    return qa_json


def get_updated_key_topics():
    # get the key topics after updates in text area
    st.session_state.list_topics = [item for item in topics_area.split("\n") if item != ""]
    st.session_state.topics_updated = True


def update_quotes_text_area():
    # getthe updated quotes after update in text area 

    # find all the quotes related to each topic 
    quotes_by_topic = re.split(r"\*{9}Topic[\d]+:.+\*{9}", quotes_text_area)
    quotes_by_topic = [x.strip() for x in quotes_by_topic if x !=""]

    # find all the topics 
    topics = re.findall(r"\*{9}(Topic[\d]+:.+)\*{9}", quotes_text_area)
    st.session_state.topics_dict = {}

    # update the topics and the quotes in the st.state_session data dictionary
    for i in range(len(topics)):
        st.session_state.topics_dict[topics[i].strip()] = {}
        quotes_list = quotes_by_topic[i].split("----------------------------------")
        st.session_state.topics_dict[topics[i].strip()]["quotes"] = [item.strip() for item in quotes_list]
        st.session_state.topics_dict[topics[i].strip()]["formated_qa"] = ""

    st.session_state.all_qa_text = ""
    # format the output 
    st.session_state.quotes_text, st.session_state.all_quotes_list = parse_response_quotes(st.session_state.topics_dict)
    st.rerun()



# Initialize session state
if 'topics_updated' not in st.session_state:
    st.session_state.topics_updated = False

if 'file_uploaded' not in st.session_state:
    st.session_state.file_uploaded = False

if 'file_content' not in st.session_state:
    st.session_state.file_content = ""

if 'keywords_extracted' not in st.session_state:
    st.session_state.keywords_extracted = False


if 'quotes_extracted' not in st.session_state:
    st.session_state.quotes_extracted = False

if 'list_topics' not in st.session_state:
    st.session_state.list_topics = []


if 'topics' not in st.session_state:
    st.session_state.topics = ""


if "topics_dict" not in st.session_state:
    st.session_state.topics_dict = {}


if 'quotes' not in st.session_state:
    st.session_state.quotes = ""


if "quotes_text" not in st.session_state:    
    st.session_state.quotes_text = ""

if "all_quotes_list" not in st.session_state:
    st.session_state.all_quotes_list = []

if "formated_qa" not in st.session_state:
    st.session_state.formated_qa = ""


if "all_qa_text" not in st.session_state:
    st.session_state.all_qa_text = ""

if "quotes_retreived" not in st.session_state:
    st.session_state.quotes_retreived = False

if "final_draft" not in st.session_state:
    st.session_state.final_draft = ""

if "memorable_quotes" not in st.session_state:
    st.session_state.memorable_quotes = []


if "custom_topics_prompt" not in st.session_state:
    st.session_state.custom_topics_prompt = ""


if "custom_qa_prompts" not in st.session_state:
    st.session_state.custom_qa_prompts = ""

st.markdown(
    """
    <style>
    .appview-container .main .block-container {
        max-width: 90%;  # Adjust the percentage as needed
        padding-top: 1rem;
        padding-right: 1rem;
        padding-left: 1rem;
        padding-bottom: 1rem;
    }

    .file-uploader-container {
        display: flex;
        justify-content: center;
        align-items: center;
    }

    .file-uploader {
        width: auto;
        max-width: 300px; /* Maximum width for the file uploader */
    }

    </style>
    """,
    unsafe_allow_html=True
)

# Apply the custom CSS
with st.container(border=False):
    st.title("📝Carlat Q/A Editor - Development Env.")
    st.caption("🚀 A streamlit editor powered by OpenAI LLM")
    st.divider()

    file = st.file_uploader(label="", type=[".docx"], key="upload", on_change = update_file_params)
    if is_file_loaded(file) and st.session_state.file_uploaded == False:
        with st.spinner("Loading in progress"):
            st.session_state.file_uploaded = True
            st.session_state.file_content = get_doc_string(file)

            #save the content of the uploaded file
            document = doc()

            # Add content to the document
            paragraphs = st.session_state.file_content.split("\n")
            for paragraph in paragraphs:
                document.add_paragraph(paragraph)

            # Save the document
            document.save("uploaded_file.docx")

            # setup vectore store 
            docs = text_splitter(st.session_state.file_content)
            docsearch = create_embedding(docs)
            docsearch.save_local('/mount/src/carlat_qa_live_update/src/vectorstore')

    custom_topics_prompts = st.checkbox(label= 'Use custom prompt to extract key topics')
    if custom_topics_prompts:
        st.session_state.custom_topics_prompt = st.text_area("Enter your custom instructions for key topics extraction")
    if  st.button("Extract key topics") and st.session_state.keywords_extracted == False:
        if(st.session_state.file_uploaded == False):
            st.warning("Please load a document first.")
        else:
            with st.spinner("Generation in progress"):
                st.session_state.topics, st.session_state.list_topics = get_key_topics(st.session_state.file_content, st.session_state.custom_topics_prompt)
                st.session_state.keywords_extracted = True
    if st.session_state.keywords_extracted == True:
        topics_area = st.text_area("Edit Key Topics", st.session_state.topics , height=300, on_change = get_updated_key_topics)
        if st.button("update key topics"):
            st.session_state.list_topics = [item for item in topics_area.split("\n") if item != ""]
            st.session_state.topics = '\n\n'.join(st.session_state.list_topics)
            st.session_state.topics_updated = True
            st.session_state.quotes_retreived = False
            st.rerun()

    if st.download_button("Download key topics", st.session_state.topics):
        if st.session_state.keywords_extracted == False:
            st.warning("Please generate key topics first.")

    if st.button("Get Quotes"): 
        if not st.session_state.keywords_extracted:
            st.warning("Please extract key topics first before getting quotes.")
        else:
            progress_text = "Operation in progress. Please wait."
            my_bar = st.progress(0, text=progress_text)
            progress = 0
            progres_increase = math.ceil(100 / len(st.session_state.list_topics))
            st.session_state.topics_dict = {}

            for (topic, i) in zip(tqdm(st.session_state.list_topics),range(len(st.session_state.list_topics))):
                progress += int(progres_increase)
                try:
                    st.session_state.topics_dict[topic] = {}
                    try:
                        st.session_state.topics_dict[topic]["quotes"] = get_quotes(topic.split(":")[1])
                    except Exception:
                        st.session_state.topics_dict[topic]["quotes"] = []
                        traceback.print_exc() 
                except Exception:
                    traceback.print_exc()

                if (progress > 100):
                    progress = 100
                my_bar.progress(progress, text=progress_text)
            my_bar.empty()

            st.session_state.btn_draft_download_status = False
            st.session_state.running = False

            # find redundant quotes to topic assingment
            redundant_quotes_dict = find_redundant_quotes(st.session_state.topics_dict)

            topics = st.session_state.topics_dict.keys()
            st.session_state.topics_dict = update_topic_assignment_all_at_once(redundant_quotes_dict, st.session_state.topics_dict, topics)
        
            # format the topics and quotes 
            st.session_state.quotes_text, st.session_state.all_quotes_list = parse_response_quotes(st.session_state.topics_dict)
            st.session_state.quotes_retreived = True
            st.session_state.topics_updated = False
            st.rerun()

    if st.session_state.quotes_retreived == True:
        quotes_text_area = st.text_area("Quotes", st.session_state.quotes_text, height=500)
    if st.button("Update quotes"):
        if st.session_state.quotes_retreived == False:
            st.warning("Please generate quotes first.")
        else:
            update_quotes_text_area()
                    
    if st.button("Highlight document"):

        if not st.session_state.quotes_retreived:
            st.warning("Please extract quotes before highlighting.")
        else:
            with st.spinner("Highlighting in progress"):
                


                document = highlight("uploaded_file.docx",st.session_state.topics_dict)

                # Read the content of the temporary file as bytes
                with open(document, "rb") as f:
                    document_bytes = f.read()

                    # Provide the bytes to the download button
                    st.download_button("Download Highlighted Document", data=document_bytes, mime="application/octet-stream", file_name="highlighted_document.docx")


    custom_qa_prompts = st.checkbox(label='Use a custom prompt to generate Q/A pairs')
    if custom_qa_prompts:
        st.session_state.custom_qa_prompts = st.text_area("Enter your custom instructions for Q/A generation")

    if st.button("Get Q/A pairs"):
        if not st.session_state.keywords_extracted:
            st.warning("Please extract key topics first before getting Q/A.")
        else:
            with st.spinner("Q/A pairs generation in progress, please wait."):
                my_bar = st.progress(0)
                progress = 0
                progres_increase = math.ceil(100 / len(st.session_state.topics_dict.keys()))
                topics = st.session_state.topics_dict.keys()
                topic_results = {topic: {"questions": [], "answers": []} for topic in topics}

                # Create a ThreadPoolExecutor with 8 threads
                with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
                    # Submit the processing of each quote for execution
                    future_to_quote = {executor.submit(process_quote,topic,quote ,st.session_state.custom_qa_prompts): (topic, quote) 
                                    for topic in topics
                                    for quote in st.session_state.topics_dict[topic]["quotes"]}
                    #Retrieve the results as they are completed
                    completed_futures, _ = concurrent.futures.wait(future_to_quote, timeout=None)

                    for future in completed_futures :
                        topic, quote = future_to_quote[future]
                        try:
                            # print("This 1 QA result per quote...")
                            qa_json = future.result()
                            st.session_state.questions = qa_json["question"]
                            st.session_state.answers = qa_json["answer"]

                            topic_results[topic]["questions"].append(qa_json["question"])
                            topic_results[topic]["answers"].append(qa_json["answer"])
                        except Exception as exc:
                            print(f"Processing exception: {exc}")

                for topic in topics:
                    st.session_state.topics_dict[topic]["formated_qa"] = format_qa_content(topic_results[topic])
                st.session_state.all_qa_text = format_qa_content_all(st.session_state.topics_dict)
                st.rerun()
   
    if st.session_state.all_qa_text !="":
        qa_text_area = st.text_area("Q/A", st.session_state.all_qa_text, height=500)

    if st.button("Generate draft", st.session_state.final_draft):
        if st.session_state.all_qa_text == "":
            st.warning("Please generate Q/A pairs first.")
        else:
            initial_draft = ""
            for topic in st.session_state.topics_dict.keys():
                initial_draft = initial_draft + st.session_state.topics_dict[topic]["formated_qa"] + "\n\n"
                    
            st.session_state.final_draft = remove_duplicates_levenshtein(initial_draft.split("\n\n"), threshold_ratio=0.65)
            # update the conversation flow and vary the questions 
            st.session_state.final_draft = make_transcript_flowful(st.session_state.topics_dict.keys(), initial_draft)

            document = doc()
            document.add_paragraph(st.session_state.final_draft)
            document.save("generated_draft.docx")

            with open("generated_draft.docx", "rb") as f:
                document_bytes = f.read()
                # Provide the bytes to the download button
                st.download_button("Download draft", data=document_bytes, mime="application/octet-stream", file_name="generated_draft.docx")
            

                # Calculate word counts for original and generated documents
                original_word_count = len(st.session_state.file_content.split())
                updated_word_count = len(st.session_state.final_draft.split())

                # Write the word counts to the Streamlit app
                st.write(f"Original Document Word Count: {original_word_count}")
                st.write(f"Generated Document Word Count: {updated_word_count}")


    if st.button("Get memorable quotes"):
        with st.spinner("Generation in progress"):
            quotes = get_memorible_quotes(st.session_state.file_content)
            st.session_state.memorable_quotes = json.loads(quotes)["quotes"]
            memorable_quotes = "\n--------------\n".join(st.session_state.memorable_quotes)

            st.text_area("Most Memorable Quotes", memorable_quotes, height=400)

        document = doc()
        document.add_paragraph(memorable_quotes)
        document.save("memorable_quotes.docx")

        with open("memorable_quotes.docx", "rb") as f:
            document_bytes = f.read()
            # Provide the bytes to the download button
            st.download_button("Download memorable quotes", data=document_bytes, mime="application/octet-stream", file_name="memorable_quotes.docx")

    if st.button("Reset"):
        st.session_state.topics_updated = False
        st.session_state.keywords_extracted = False
        st.session_state.quotes_extracted = False
        st.session_state.list_topics = []
        st.session_state.topics = ""
        st.session_state.topics_dict = {}
        st.session_state.quotes = ""
        st.session_state.quotes_text = ""
        st.session_state.all_quotes_list = []
        st.session_state.formated_qa = ""
        st.session_state.all_qa_text = ""
        st.session_state.quotes_retreived = False
        st.session_state.final_draft = ""
        st.session_state.memorable_quotes = []
        st.empty()
        st.rerun()
